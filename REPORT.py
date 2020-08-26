import pandas as pd
import os
import time
from mailmerge import MailMerge
from datetime import date, datetime
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


'''
The idea for writing this is after reading the AUTOMATE THE BORING STUFF book.
I am an HR manager and my job partly consists of preparing a wide variety of reports for management, government and statistics.
I've already become a pro in Excel in VBA, but I really like programming and the Python language.
So I decided to automate a bit of the boring job of reporting my department's performance to senior management.
Previously, you had to do this:
Due to the fact that I do not have direct access to the database, at first I had to generate 
a ready-made report (xls file (the time of formation by the program is 1 minute) and then work with it 
(but I have an idea and understanding of how to write the required SQL query , to unload this information from the database), 
open data on people, make pivot tables, remove unnecessary things, copy to Word (do not ask why))) and then print and sign. 
On average, if not distracted, it took about 20 minutes.

Now life has become much easier and better))). 
I got new knowledge, skills and abilities. 
And I get the finished report in 0.5 seconds, taking into account the time to save the original xls file - 1 minute and 0.5 seconds.


P.S. I'm understand that junior's level and this file not all in PEP rules. 
But i'm write his file with maximum readability to whom who will want repeat this in their work
and of course automete the boring stuff  :)
'''

 
start = time.time()

path = "D:\\REPORTS\\REPORT1\\"

sex = ["ж", "м"]
category = [1, 2, 3, 4, 5]


def df(name):
    df = pd.read_excel("D:\\REPORTS\\REPORT1\\REPORT1.xls", index_col=0, sheet_name=name)  # xls file and sheet with all emploees data

    average_age = 0
    taday_date = 0
    month = 0
    fired_reason = 0
    hired_from = 0

    if name == "sheet1":  # all employees
        df['year'] = pd.DatetimeIndex(df['d_rogden']).year
        df['age'] = 2020 - df['year']
        average_age = round(df['age'].mean())
        taday_date = datetime.today()
        month = str(taday_date.month - 1)

    if name == "sheet2":  # hired in report period
        df = df[(df["p_priem"] != 9) & (df["p_priem"] != 17)]
        hired_from = df["namepriem"].value_counts().to_dict()
        if 'towards organs' in hired_from.keys():
            hired_from['towards government'] = hired_from.pop('towards organs')
            hired_from = {k: v for k, v in sorted(hired_from.items(), key=lambda item: item[1], reverse=True)}

    if name == "sheet3":  #fired in report period
        df = df[(df["_priem"] != 17) & (df["nameyvol"] != "transfer to №1")]
        fired_reason = df["nameyvol"].value_counts().to_dict()

    count_employee_category = df["kkat"].value_counts().to_dict()
    person_category_dict = {k: v for k, v in sorted(count_employee_category.items(), key=lambda item: item[0])}
    sex_all = df["pol"].value_counts().to_dict()
    sex_women = {k: v for k, v in sorted(sex_all.items(), key=lambda item: item[0])}

    # add zero to key : value pairs if some category or sex doesn't hired/fired in report period
    for i in sex:
        if i not in sex_women.keys():
            sex_women.update({i: 0})
    for i in category:
        if i not in person_category_dict.keys():
            person_category_dict.update({i: 0})

    return [person_category_dict, sex_women, average_age, month, df, hired_from, fired_reason]


person_category_dict, sex_women, average_age, month, _, _, _ = df("sheet1")
hired_category_dict, hired_sex, _, _, _, hired_from, _ = df("sheet2")
fired_category_dict, fired_sex, _, _, _, _, fired_reason = df("sheet3")


template_1 = "D:\\REPORTS\\REPORT1\\blank.docx"

# here i'm count last day in report period
today = date.today()
last_day = date(today.year, today.month, 1) - relativedelta(days=1)
date = last_day.strftime('%d.%m.%Y')

# filling template docx
document = MailMerge(template_1)
document.merge(
    all_emp=str(sum(person_category_dict.values())),
    all_itr=str(person_category_dict[3] + person_category_dict[4] + person_category_dict[5]),
    all_ruk=str(person_category_dict[3]),
    all_spec=str(person_category_dict[5]),
    all_drsl=str(person_category_dict[4]),
    all_rab=str(person_category_dict[1] + person_category_dict[2]),
    all_women=str(sex_women["ж"]),
    date=str(date),
    average_age=str(average_age),
    month=str(month),

    all_hired=str(sum(hired_category_dict.values())),
    hired_ruk=str(hired_category_dict[3]),
    hired_spec=str(hired_category_dict[5]),
    hired_drsl=str(hired_category_dict[4]),
    hired_rab1=str(hired_category_dict[1] + hired_category_dict[2]),
    hired_women=str(hired_sex["ж"]),
    hired_men=str(hired_sex["м"]),

    all_fired=str(sum(fired_category_dict.values())),
    fired_ruk=str(fired_category_dict[3]),
    fired_spec=str(fired_category_dict[5]),
    fired_drsl=str(fired_category_dict[4]),
    fired_rab=str(fired_category_dict[1] + fired_category_dict[2]),
    fired_women=str(fired_sex["ж"]),
    fired_men=str(fired_sex["м"])
)

filename = "Report (working with stuff) for " + month + " month 2020.docx"
document.write(path + filename)  # save file to folder
print(filename)

# here we create 2 tables, one below the other.
# The first table contains the names of the fields, the second table contains summary data by type of hired and fired reasons.


def table(name):
    # customizing first table
    word_document = Document(path + filename)
    table0 = word_document.add_table(0, 0)
    table0.style = word_document.styles["Table Grid"]
    first_column_width = 15
    second_column_with = 2.5
    table0.add_column(Cm(first_column_width))
    table0.add_column(Cm(second_column_with))
    table0.add_row()
    header_cells = table0.rows[-1].cells
    if name == hired_from:
        header_cells[0].text = "Hired type"
    else:
        header_cells[0].text = "Fired reason"
    header_cells[1].text = "employee"
    table0.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    table0.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True
    table0.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table0.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # customizing second table
    table1 = word_document.add_table(0, 0)  # we add rows iteratively
    table1.style = word_document.styles["Table Grid"]
    first_column_width = 15
    second_column_with = 2.5
    table1.add_column(Cm(first_column_width))
    table1.add_column(Cm(second_column_with))

    for index, stat_item in enumerate(name.items()):
        table1.add_row()
        stat_name, stat_result = stat_item
        row = table1.rows[index]
        row.cells[0].text = str(stat_name)
        row.cells[1].text = str(stat_result)
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    word_document.add_paragraph()
    word_document.save(path + filename)


table(hired_from)
table(fired_reason)

end = time.time()
print(end - start)

input("\nPress any key to exit...")

os.startfile(path)
